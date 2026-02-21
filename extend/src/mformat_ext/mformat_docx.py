#! /usr/local/bin/python3
"""Extension of the MultiFormat class for DOCX files."""

# Copyright (c) 2025 - 2026 Tom Björkholm
# MIT License
#

from typing import Optional, Callable, Any
from docx import Document
from docx.document import Document as DocumentObject
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from mformat.mformat import FormatterDescriptor, MultiFormat
from mformat.mformat_state import MultiFormatState, Formatting

_MAX_LIST_LEVEL = 5
"""Maximum supported list nesting level for DOCX format."""


class MultiFormatDocx(MultiFormat):
    """Extension of the MultiFormat class for DOCX files."""

    def __init__(self, file_name: str, url_as_text: bool = False,
                 file_exists_callback: Optional[Callable[[str], None]] = None):
        """Initialize the MultiFormatDocx class.

        Args:
            file_name: The name of the file to write to.
            url_as_text: Format URLs as text not clickable URLs.
            file_exists_callback: A callback function to call if the file
                                  already exists. Return to allow the file to
                                  be overwritten. Raise an exception to prevent
                                  the file from being overwritten.
                                  (May for instance save existing file as
                                  backup.)
                                  (Default is to raise an exception.)
        """
        self.doc: DocumentObject = Document()
        self.current_paragraph: Optional[Paragraph] = None
        self._current_abstract_num: Any = None
        self._current_num_id: int = -1
        super().__init__(file_name=file_name, url_as_text=url_as_text,
                         file_exists_callback=file_exists_callback)

    @classmethod
    def file_name_extension(cls) -> str:
        """Get the file name extension for the formatter."""
        return '.docx'

    @classmethod
    def get_arg_desciption(cls) -> FormatterDescriptor:
        """Get the description of the arguments for the formatter."""
        return FormatterDescriptor(name='docx', mandatory_args=[],
                                   optional_args=[])

    def open(self) -> None:
        """Open the file.

        Avoid using this method directly.
        Use as a context manager instead, using a with statement.
        """
        self.doc.save(self.file_name)

    def _close(self) -> None:
        """Close the file.

        Avoid using this method directly.
        Use as a context manager instead, using a with statement.
        """
        if self.state == MultiFormatState.EMPTY:
            return
        self.doc.save(self.file_name)

    def _write_file_prefix(self) -> None:
        """Write the file prefix.

        For DOCX files, this is a no-op since the document
        structure is handled by python-docx.
        """

    def _write_file_suffix(self) -> None:
        """Write the file suffix.

        For DOCX files, this is a no-op since the document
        structure is handled by python-docx.
        """

    def _start_paragraph(self) -> None:
        """Start a paragraph."""
        self.current_paragraph = self.doc.add_paragraph()

    def _end_paragraph(self) -> None:
        """End a paragraph."""
        self.current_paragraph = None

    def _start_heading(self, level: int) -> None:
        """Start a heading.

        Args:
            level: The level of the heading (1-9).
        """
        self.current_paragraph = self.doc.add_heading(level=level)

    def _end_heading(self, level: int) -> None:
        """End a heading.

        Args:
            level: The level of the heading (1-9).
        """
        self.current_paragraph = None

    def _write_text(self, text: str, state: MultiFormatState,
                    formatting: Formatting) -> None:
        """Write text into current item (paragraph, bullet list item, etc.).

        Args:
            text: The text to write into the current item.
            state: The state of the current item.
            formatting: The formatting of the text.
        """
        if self.current_paragraph is None:  # pragma: no cover # noqa: E501
            raise RuntimeError('No current paragraph to write text into')
        run = self.current_paragraph.add_run(text)
        if formatting.bold:
            run.bold = True
        if formatting.italic:
            run.italic = True

    def _write_url(self, url: str, text: Optional[str],
                   state: MultiFormatState,
                   formatting: Formatting) -> None:
        """Write a URL into current item (paragraph, bullet list item, etc.).

        Args:
            url: The URL to write into the current item.
            text: The text to display for the URL.
            state: The state of the current item.
            formatting: The formatting of the text.
        """  # pylint: disable=too-many-arguments,too-many-positional-arguments
        if self.current_paragraph is None:  # pragma: no cover # noqa: E501
            raise RuntimeError('No current paragraph to write URL into')
        if not text:
            text = url
        self._add_hyperlink(self.current_paragraph, url, text, formatting)

    def _add_hyperlink(self, paragraph: Paragraph, url: str,
                       text: str, formatting: Formatting) -> None:
        """Add a clickable hyperlink to a paragraph.

        Args:
            paragraph: The paragraph to add the hyperlink to.
            url: The URL for the hyperlink.
            text: The display text for the hyperlink.
            formatting: The formatting (bold/italic) for the text.
        """
        # Create relationship for the hyperlink
        # The relationship type for external hyperlinks
        rel_type = ('http://schemas.openxmlformats.org/officeDocument/'
                    '2006/relationships/hyperlink')
        r_id = self.doc.part.relate_to(url, rel_type, is_external=True)
        # Create hyperlink element
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        # Create run inside hyperlink
        run_element = OxmlElement('w:r')
        run_props = OxmlElement('w:rPr')
        # Add blue color for link appearance
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        run_props.append(color)
        # Add underline for link appearance
        underline = OxmlElement('w:u')
        underline.set(qn('w:val'), 'single')
        run_props.append(underline)
        # Add bold if requested
        if formatting.bold:
            bold = OxmlElement('w:b')
            run_props.append(bold)
        # Add italic if requested
        if formatting.italic:
            italic = OxmlElement('w:i')
            run_props.append(italic)
        run_element.append(run_props)
        # Add text element
        text_element = OxmlElement('w:t')
        text_element.text = text
        run_element.append(text_element)
        hyperlink.append(run_element)
        # Append hyperlink to paragraph
        # pylint: disable=protected-access
        paragraph._p.append(hyperlink)

    def _write_code_in_text(self, text: str,
                            state: MultiFormatState) -> None:
        """Write code in text.

        Write code text into the current paragraph, heading, bullet list item
        or numbered point list item.

        Args:
            text: The text to add to the code block.
            state: The state of the current item.
        """
        assert state in [MultiFormatState.BULLET_LIST_ITEM,
                         MultiFormatState.NUMBERED_LIST_ITEM,
                         MultiFormatState.HEADING,
                         MultiFormatState.PARAGRAPH,
                         MultiFormatState.CODE_BLOCK]
        assert isinstance(text, str)
        if self.current_paragraph is None:  # pragma: no cover # noqa: E501
            raise RuntimeError('No current paragraph to write code into')
        run = self.current_paragraph.add_run(text)
        # Set monospace font
        run.font.name = 'Courier New'

    # ========================================================
    # List numbering infrastructure
    # ========================================================

    @staticmethod
    def _validate_list_depth(level: int) -> None:
        """Validate that the list level is within DOCX limits.

        Args:
            level: The list nesting level to validate.
        Raises:
            RuntimeError: If level exceeds the maximum.
        """
        if level > _MAX_LIST_LEVEL:
            raise RuntimeError(
                f'DOCX format supports at most '
                f'{_MAX_LIST_LEVEL} nesting levels for '
                f'lists, but level {level} was requested.')

    def _get_numbering_xml(self) -> Any:
        """Get the numbering XML root element.

        Returns:
            The w:numbering XML element (CT_Numbering).
        """
        # pylint: disable=protected-access
        return (
            self.doc.part
            .numbering_part._element
        )

    def _next_abstract_num_id(self) -> int:
        """Get the next available abstractNumId.

        Returns:
            The next available abstractNumId value.
        """
        numbering = self._get_numbering_xml()
        ids = [
            int(an.get(qn('w:abstractNumId')))
            for an in numbering.findall(
                qn('w:abstractNum'))
        ]
        return max(ids) + 1 if ids else 0

    @staticmethod
    def _create_level_xml(
            ilvl: int, bullet: bool) -> Any:
        """Create a numbering level definition element.

        Args:
            ilvl: The indent level (0-based).
            bullet: True for bullet, False for decimal.
        Returns:
            The w:lvl XML element.
        """
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), str(ilvl))
        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)
        num_fmt = OxmlElement('w:numFmt')
        fmt = 'bullet' if bullet else 'decimal'
        num_fmt.set(qn('w:val'), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement('w:lvlText')
        if bullet:
            lvl_text.set(qn('w:val'), '\u2022')
        else:
            lvl_text.set(
                qn('w:val'), f'%{ilvl + 1}.')
        lvl.append(lvl_text)
        lvl_jc = OxmlElement('w:lvlJc')
        lvl_jc.set(qn('w:val'), 'left')
        lvl.append(lvl_jc)
        p_pr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        ind.set(
            qn('w:left'), str(720 * (ilvl + 1)))
        ind.set(qn('w:hanging'), '360')
        p_pr.append(ind)
        lvl.append(p_pr)
        if bullet:
            r_pr = OxmlElement('w:rPr')
            r_fonts = OxmlElement('w:rFonts')
            r_fonts.set(qn('w:ascii'), 'Symbol')
            r_fonts.set(qn('w:hAnsi'), 'Symbol')
            r_fonts.set(qn('w:hint'), 'default')
            r_pr.append(r_fonts)
            lvl.append(r_pr)
        return lvl

    def _create_abstract_num(
            self, bullet: bool) -> Any:
        """Create a multi-level abstract numbering definition.

        All five levels are initialized with the same
        format (bullet or decimal). Individual levels can
        be changed later via _update_level_format for mixed
        lists.

        Args:
            bullet: True for bullets, False for decimal.
        Returns:
            The w:abstractNum XML element.
        """
        numbering = self._get_numbering_xml()
        abstract_id = self._next_abstract_num_id()
        abstract_num = OxmlElement('w:abstractNum')
        abstract_num.set(
            qn('w:abstractNumId'), str(abstract_id))
        multi = OxmlElement('w:multiLevelType')
        multi.set(
            qn('w:val'), 'hybridMultilevel')
        abstract_num.append(multi)
        for i in range(_MAX_LIST_LEVEL):
            abstract_num.append(
                self._create_level_xml(i, bullet))
        first_num = numbering.find(qn('w:num'))
        if first_num is not None:
            first_num.addprevious(abstract_num)
        else:
            numbering.append(abstract_num)
        return abstract_num

    def _create_num_instance(
            self,
            abstract_num: Any) -> int:
        """Create a numbering instance for an abstract def.

        Args:
            abstract_num: The w:abstractNum element.
        Returns:
            The numId of the new numbering instance.
        """
        numbering = self._get_numbering_xml()
        abstract_id = int(
            abstract_num.get(qn('w:abstractNumId')))
        ct_num = numbering.add_num(abstract_id)
        return ct_num.numId  # type: ignore[no-any-return]

    def _update_level_format(
            self, ilvl: int, bullet: bool) -> None:
        """Update format of a level in current abstractNum.

        Used when a nested list has a different type than
        its parent (e.g., bullets nested in numbered).

        Args:
            ilvl: The indent level to update (0-based).
            bullet: True for bullet, False for decimal.
        """
        if self._current_abstract_num is None:
            return  # pragma: no cover
        for lvl in self._current_abstract_num.findall(
                qn('w:lvl')):
            if lvl.get(qn('w:ilvl')) == str(ilvl):
                idx = list(
                    self._current_abstract_num
                ).index(lvl)
                self._current_abstract_num.remove(lvl)
                new_lvl = self._create_level_xml(
                    ilvl, bullet)
                self._current_abstract_num.insert(
                    idx, new_lvl)
                return

    def _set_paragraph_list_props(
            self, ilvl: int) -> None:
        """Set numbering properties on current paragraph.

        Args:
            ilvl: The indent level (0-based).
        """
        if self.current_paragraph is None:  # pragma: no cover # noqa: E501
            raise RuntimeError(
                'No current paragraph for list props')
        # pylint: disable=protected-access
        p_pr = (
            self.current_paragraph._p.get_or_add_pPr()
        )
        num_pr = OxmlElement('w:numPr')
        ilvl_elem = OxmlElement('w:ilvl')
        ilvl_elem.set(qn('w:val'), str(ilvl))
        num_pr.append(ilvl_elem)
        num_id_elem = OxmlElement('w:numId')
        num_id_elem.set(
            qn('w:val'), str(self._current_num_id))
        num_pr.append(num_id_elem)
        p_pr.append(num_pr)

    def _start_list(
            self, level: int, bullet: bool) -> None:
        """Start or continue a list group.

        Creates new numbering for top-level lists. For
        nested lists, updates the level format if the type
        differs from the initial list type.

        Args:
            level: The list nesting level (1-based).
            bullet: True for bullet, False for numbered.
        """
        self._validate_list_depth(level)
        if level == 1:
            abstract_num = self._create_abstract_num(
                bullet)
            self._current_abstract_num = abstract_num
            self._current_num_id = (
                self._create_num_instance(abstract_num))
        else:
            self._update_level_format(
                level - 1, bullet)

    def _end_list(self, level: int) -> None:
        """End a list level, reset state at top level.

        Args:
            level: The list nesting level (1-based).
        """
        if level == 1:
            self._current_abstract_num = None
            self._current_num_id = -1

    # ========================================================
    # Bullet list methods
    # ========================================================

    def _start_bullet_list(self, level: int) -> None:
        """Start a bullet list.

        Args:
            level: The level of the bullet list (1-5).
        """
        assert isinstance(level, int)
        self._start_list(level, bullet=True)

    def _end_bullet_list(self, level: int) -> None:
        """End a bullet list.

        Args:
            level: The level of the bullet list (1-5).
        """
        assert isinstance(level, int)
        self._end_list(level)

    def _start_bullet_item(self, level: int) -> None:
        """Start a bullet item.

        Args:
            level: The level of the bullet item (1-5).
        """
        assert isinstance(level, int)
        self._validate_list_depth(level)
        self.current_paragraph = (
            self.doc.add_paragraph())
        self._set_paragraph_list_props(
            ilvl=level - 1)

    def _end_bullet_item(self, level: int) -> None:
        """End a bullet item.

        Args:
            level: The level of the bullet item (1-5).
        """
        assert isinstance(level, int)
        self.current_paragraph = None

    # ========================================================
    # Numbered list methods
    # ========================================================

    def _start_numbered_list(self, level: int) -> None:
        """Start a numbered list.

        Args:
            level: The level of the numbered list (1-5).
        """
        assert isinstance(level, int)
        self._start_list(level, bullet=False)

    def _end_numbered_list(self, level: int) -> None:
        """End a numbered list.

        Args:
            level: The level of the numbered list (1-5).
        """
        assert isinstance(level, int)
        self._end_list(level)

    def _start_numbered_item(self, level: int, num: int,
                             full_number: str) -> None:
        """Start a numbered item.

        Word handles the numbering automatically through
        the numbering definitions. The num and full_number
        parameters are not used for DOCX output.

        Args:
            level: The level of the numbered item (1-5).
            num: The number of the item (unused in DOCX).
            full_number: The full hierarchical number
                         (unused in DOCX).
        """
        assert isinstance(level, int)
        assert isinstance(num, int)
        assert isinstance(full_number, str)
        self._validate_list_depth(level)
        self.current_paragraph = (
            self.doc.add_paragraph())
        self._set_paragraph_list_props(
            ilvl=level - 1)

    def _end_numbered_item(self, level: int,
                           num: int) -> None:
        """End a numbered item.

        Args:
            level: The level of the numbered item (1-5).
            num: The number of the item.
        """
        assert isinstance(level, int)
        assert isinstance(num, int)
        self.current_paragraph = None

    def _start_table(self, num_columns: int) -> None:
        """Start a table.

        Args:
            num_columns: The number of columns in the table.
        """
        assert isinstance(num_columns, int)
        # Add an empty paragraph before the table for spacing
        # This separates the table from previous content
        self.doc.add_paragraph()
        # Note: Actual table creation happens in _write_table_first_row

    def _end_table(self, num_columns: int, num_rows: int) -> None:
        """End a table.

        Args:
            num_columns: The number of columns in the table.
            num_rows: The number of rows in the table.
        """
        assert isinstance(num_columns, int)
        assert isinstance(num_rows, int)
        # Add an empty paragraph after the table for spacing
        self.doc.add_paragraph()

    def _write_table_first_row(self, first_row: list[str],
                               formatting: Formatting) -> None:
        """Write the first row of a table.

        Args:
            first_row: The first row of the table.
            formatting: The formatting of the text in each cell.
        """
        assert isinstance(first_row, list)
        assert isinstance(formatting, Formatting)
        # Create the table with the first row
        table = self.doc.add_table(rows=1, cols=len(first_row))
        table.style = 'Table Grid'
        # Fill in the first row
        for idx, cell_text in enumerate(first_row):
            cell = table.rows[0].cells[idx]
            para = cell.paragraphs[0]
            run = para.add_run(cell_text)
            if formatting.bold:
                run.bold = True
            if formatting.italic:
                run.italic = True

    def _write_table_row(self, row: list[str], formatting: Formatting,
                         row_number: int) -> None:
        """Write a row of a table.

        Args:
            row: The row to add to the table.
            formatting: The formatting of the text in each cell.
            row_number: The row number (0-based).
        """
        assert isinstance(row, list)
        assert isinstance(formatting, Formatting)
        assert isinstance(row_number, int)
        # Get the last table in the document
        table = self.doc.tables[-1]
        # Add a new row
        new_row = table.add_row()  # type: ignore[no-untyped-call]
        # Fill in the row
        for idx, cell_text in enumerate(row):
            cell = new_row.cells[idx]
            para = cell.paragraphs[0]
            run = para.add_run(cell_text)
            if formatting.bold:
                run.bold = True
            if formatting.italic:
                run.italic = True

    def _start_code_block(self, programming_language: Optional[str]) -> None:
        """Start a code block.

        Args:
            programming_language: The programming language of the code block.
        """
        assert programming_language is None or \
            isinstance(programming_language, str)
        # Create a paragraph with a code/verbatim style
        # python-docx doesn't have a built-in code style,
        # so we'll use 'No Spacing' style and monospace font
        self.current_paragraph = self.doc.add_paragraph()
        self.current_paragraph.style = 'No Spacing'

    def _end_code_block(self, programming_language: Optional[str]) -> None:
        """End a code block.

        Args:
            programming_language: The programming language of the code block.
        """
        assert programming_language is None or \
            isinstance(programming_language, str)
        self.current_paragraph = None

    def _write_code_block(self, text: str,
                          programming_language: Optional[str]) -> None:
        """Write a code block.

        Args:
            text: The text to add to the code block.
            programming_language: The programming language of the code block.
        """
        # _start_code_block and _end_code_block handle the BLOCK aspects,
        # so we only need to write the text inside the code block.
        assert programming_language is None or \
            isinstance(programming_language, str)
        self._write_code_in_text(text, state=MultiFormatState.CODE_BLOCK)

    def _encode_text(self, text: str) -> str:
        """Encode text (escape special characters)."""
        # No encoding needed for DOCX
        return text
